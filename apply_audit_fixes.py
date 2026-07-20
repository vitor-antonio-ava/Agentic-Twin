"""
apply_audit_fixes.py
Applies the Doc-Audit findings from MyMeetingsTwin_Doc_Audit.html
to every file in OneDrive_1_7-17-2026/.

Run:  .venv\Scripts\python.exe apply_audit_fixes.py
"""

import shutil, re, copy, os
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl

FOLDER = Path(r"c:\Dev\GSK - Agentic Twin\OneDrive_1_7-17-2026")
BACKUP = Path(r"c:\Dev\GSK - Agentic Twin\OneDrive_1_7-17-2026_backup")

# ── helpers ──────────────────────────────────────────────────────────────────

def backup_folder():
    if not BACKUP.exists():
        shutil.copytree(FOLDER, BACKUP)
        print(f"  [backup] {BACKUP}")
    else:
        print(f"  [backup] already exists, skipping")

def run_in_paragraph(para, old: str, new: str) -> bool:
    """Replace text across runs in a paragraph, preserving formatting of first run."""
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    return True

def replace_in_doc(doc: Document, old: str, new: str) -> int:
    """Replace text everywhere in a document (paragraphs + table cells)."""
    count = 0
    for para in doc.paragraphs:
        if run_in_paragraph(para, old, new):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if run_in_paragraph(para, old, new):
                        count += 1
    return count

def find_paragraphs_containing(doc: Document, text: str):
    results = []
    for i, para in enumerate(doc.paragraphs):
        if text in para.text:
            results.append((i, para))
    return results

def delete_paragraph(para):
    """Remove a paragraph from the document."""
    p = para._element
    p.getparent().remove(p)

def delete_table_row_containing(doc: Document, text: str) -> bool:
    """Delete first table row whose text contains the given string."""
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(c.text for c in row.cells)
            if text in row_text:
                tbl = table._tbl
                tr = row._tr
                tbl.remove(tr)
                return True
    return False

def add_note_paragraph(doc: Document, after_para_idx: int, note_text: str, style: str = "Normal"):
    """Insert a bold note paragraph after the given paragraph index."""
    # We add it by appending to the parent – simple approach: add after the element
    para = doc.paragraphs[after_para_idx]
    new_para = OxmlElement("w:p")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    b = OxmlElement("w:b")
    rPr.append(b)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = note_text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    new_para.append(r)
    para._element.addnext(new_para)

def save_doc(doc: Document, path: Path, label: str):
    doc.save(path)
    print(f"  [saved] {label}")

# ── DIA definition (universal fix) ───────────────────────────────────────────
DIA_OLD = "Data Input & Analysis"
DIA_NEW = "Decision / Input / Awareness"

# ── 1. TDD.docx ──────────────────────────────────────────────────────────────
def fix_tdd():
    path = FOLDER / "My Meetings Twin - TDD.docx"
    doc = Document(path)
    changes = []

    # (a) DIA definition everywhere
    n = replace_in_doc(doc, DIA_OLD, DIA_NEW)
    changes.append(f"DIA definition ×{n}")

    # (b) Fix §3.6: "Initiate DIA" → "Start Meeting Configuration"
    n = replace_in_doc(doc, "Initiate DIA", "Start Meeting Configuration")
    if n:
        changes.append(f"'Initiate DIA' → 'Start Meeting Configuration' ×{n}")

    # (c) Fix §3.7 GetUserConfiguredMeetings description
    n = replace_in_doc(doc,
        "Retrieves all the User Teams Meetings where the current user is the organizer for",
        "Retrieves the user's previously configured meetings stored in Dataverse (flow: GetUserMeetingConfigurationsV2)")
    if n:
        changes.append("Fixed §3.7 GetUserConfiguredMeetings description")

    # (d) Delete the RPA/Automation Anywhere/OLEDB paragraph (§3.28)
    rpa_marker = "This process is built based on AA tasks and scripts"
    deleted = False
    for para in list(doc.paragraphs):
        if rpa_marker in para.text:
            delete_paragraph(para)
            deleted = True
            changes.append("Deleted RPA/OLEDB paragraph (§3.28)")
            break

    # (e) Fix the self-contradiction on scheduling: replace the "No background processing" / 
    #     "not applicable" claims with the correct model.
    sched_replacements = [
        ("No background processing", "Background processing: an hourly recurrence flow (PA_SendDIANotificationToOrganizer) runs unattended; four Dataverse-webhook flows fire on row changes"),
        ("No unattended or background execution", "Unattended execution exists: PA_SendDIANotificationToOrganizer runs on a Recurrence trigger every 1 hour; event-driven Dataverse flows notify organizers/owners on row changes"),
        ("Scheduling: Not applicable", "Scheduling: Hourly recurrence (PA_SendDIANotificationToOrganizer) plus event-driven Dataverse webhook flows; the conversational agent itself is user-initiated"),
        ("does not run on a schedule", "runs an hourly notification schedule (PA_SendDIANotificationToOrganizer) in addition to user-initiated conversations"),
        ("will essentially run daily to check for notifications", "runs an hourly recurrence flow (every 1 hour) to check for and send DIA preparation notifications"),
    ]
    for old, new in sched_replacements:
        n = replace_in_doc(doc, old, new)
        if n:
            changes.append(f"Scheduling fix: '{old[:40]}…' ×{n}")

    # (f) Fix §2.4 Acronyms table entry
    n = replace_in_doc(doc, "Data Input & Analysis", "Decision / Input / Awareness")
    # already handled above, this is a no-op if already replaced

    save_doc(doc, path, "TDD.docx")
    print(f"    changes: {', '.join(changes) if changes else 'none'}")

# ── 2. AI Accountability Report.xlsx ─────────────────────────────────────────
def fix_ai_accountability():
    path = FOLDER / "My Meetings Twin - AI Accountability Report.xlsx"
    wb = openpyxl.load_workbook(path)
    changes = []

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    if DIA_OLD in cell.value:
                        cell.value = cell.value.replace(DIA_OLD, DIA_NEW)
                        changes.append(f"DIA def in [{ws.title}]!{cell.coordinate}")
                    # Fix "does not involve custom-developed models" note
                    if "does not involve custom-developed models" in cell.value:
                        cell.value = cell.value + " [AUDIT NOTE: Solution includes 7 active AI Builder prompt models (RedKPI-Interpreter, Draft-DIA, Transcript-Analysis, Calc-Attendees, HTML-Slides, KPI-Selection, MSProject-Selection) and 13 orphaned models. These must be disclosed.]"
                        changes.append(f"AI models disclosure note in [{ws.title}]!{cell.coordinate}")
                    if "does not trigger actions" in cell.value or "does not perform actions" in cell.value:
                        cell.value = cell.value + " [AUDIT NOTE: Agent DOES create follow-up meetings & sends invites, sends emails, and creates Planner tasks. Enumerate write actions and revise Model Influence / Decision Consequence ratings.]"
                        changes.append(f"Write-actions note in [{ws.title}]!{cell.coordinate}")
                    if ("does not run unattended" in cell.value or 
                        "user-triggered and does not run unattended" in cell.value):
                        cell.value = cell.value + " [AUDIT NOTE: Unattended flows exist — PA_SendDIANotificationToOrganizer runs hourly; DIANotificationtoTaskOwnersandAttendees fires on DIA task creation. Revise ratings accordingly.]"
                        changes.append(f"Unattended-flows note in [{ws.title}]!{cell.coordinate}")

    wb.save(path)
    print(f"  [saved] AI Accountability Report.xlsx")
    print(f"    changes: {', '.join(changes) if changes else 'none (check manually — DIA text may be in merged cells or formatted differently)'}")

# ── 3. User Requirements.docx ────────────────────────────────────────────────
def fix_user_requirements():
    path = FOLDER / "My Meetings Twin - User Requirements.docx"
    doc = Document(path)
    changes = []

    # (a) DIA definition
    n = replace_in_doc(doc, DIA_OLD, DIA_NEW)
    changes.append(f"DIA definition ×{n}")

    # (b) Interface requirement: "Microsoft Graph" → add note about missing connectors
    n = replace_in_doc(doc,
        "Microsoft Graph",
        "Microsoft Graph (Note: actual connectors also include Planner, Microsoft Teams, Office 365 Outlook, and SharePoint Online — add UR-07xx requirements for each)")
    if n:
        changes.append(f"Interface/connector note ×{n}")

    # (c) Fix availability requirement
    n = replace_in_doc(doc,
        "shall remain available during standard business hours",
        "shall remain available during standard business hours. Note: the automated notification processing (PA_SendDIANotificationToOrganizer) runs continuously on an hourly schedule, independent of business hours.")
    if n:
        changes.append("Availability requirement updated")

    save_doc(doc, path, "User Requirements.docx")
    print(f"    changes: {', '.join(changes) if changes else 'none'}")

# ── 4. EUD Checklist.docx ────────────────────────────────────────────────────
def fix_eud_checklist():
    path = FOLDER / "My Meetings Twin - EUD Checklist.docx"
    doc = Document(path)
    changes = []

    # (a) DIA definition
    n = replace_in_doc(doc, DIA_OLD, DIA_NEW)
    changes.append(f"DIA definition ×{n}")

    # (b) Connector list expansion: add missing connectors
    connector_pairs = [
        ("Graph, Dataverse, Power BI, Planner",
         "Dataverse, Office 365 Outlook, Office 365 Users, Office 365 Groups, Planner, Microsoft Teams, Azure AD, SharePoint Online, Power BI"),
        ("Graph, Dataverse, Power BI",
         "Dataverse, Office 365 Outlook, Office 365 Users, Office 365 Groups, Planner, Microsoft Teams, Azure AD, SharePoint Online, Power BI"),
        ("Microsoft Graph, Dataverse, Power BI, Planner",
         "Dataverse, Office 365 Outlook, Office 365 Users, Office 365 Groups, Planner, Microsoft Teams, Azure AD, SharePoint Online, Power BI"),
        ("Microsoft Graph, Dataverse, Power BI",
         "Dataverse, Office 365 Outlook, Office 365 Users, Office 365 Groups, Planner, Microsoft Teams, Azure AD, SharePoint Online, Power BI"),
    ]
    for old_c, new_c in connector_pairs:
        n = replace_in_doc(doc, old_c, new_c)
        if n:
            changes.append(f"Connector list expanded ×{n}")
            break

    # (c) Fix Solution Version
    n = replace_in_doc(doc, "Solution Version '1.0.0'", "Solution Version 1.0.0.7")
    if n == 0:
        n = replace_in_doc(doc, "1.0.0", "1.0.0.7 (deployed build)")
        if n:
            changes.append("Solution version updated to 1.0.0.7")

    save_doc(doc, path, "EUD Checklist.docx")
    print(f"    changes: {', '.join(changes) if changes else 'none'}")

# ── 5. GSC RAI Process.docx ──────────────────────────────────────────────────
def fix_gsc_rai_process():
    path = FOLDER / "My Meetings Twin - GSC RAI Process.docx"
    doc = Document(path)
    changes = []

    # (a) DIA definition
    n = replace_in_doc(doc, DIA_OLD, DIA_NEW)
    changes.append(f"DIA definition ×{n}")

    # (b) Data source list: add missing sources
    data_src_pairs = [
        ("Microsoft Graph, Dataverse, Power BI, user inputs",
         "Dataverse, Power BI, Microsoft Teams (transcripts), Office 365 Outlook (email/calendar), Planner (tasks), SharePoint Online (deck storage), user inputs"),
        ("Microsoft Graph, Dataverse, Power BI",
         "Dataverse, Power BI, Microsoft Teams (transcripts), Office 365 Outlook (email/calendar), Planner (tasks), SharePoint Online (deck storage)"),
    ]
    for old_d, new_d in data_src_pairs:
        n = replace_in_doc(doc, old_d, new_d)
        if n:
            changes.append(f"Data sources list updated ×{n}")
            break

    # (c) AI models disclosure note
    n = replace_in_doc(doc,
        "conversational assistant",
        "conversational assistant (Note: solution also uses 7 active AI Builder prompt models including RedKPI-Interpreter, Draft-DIA, and Transcript-Analysis — these must be disclosed in the AI intake)")
    if n:
        changes.append("AI models disclosure note added")

    save_doc(doc, path, "GSC RAI Process.docx")
    print(f"    changes: {', '.join(changes) if changes else 'none'}")

# ── 6. BCP.docx ──────────────────────────────────────────────────────────────
def fix_bcp():
    path = FOLDER / "My Meetings Twin - BCP.docx"
    doc = Document(path)
    changes = []

    # (a) Delete the SAP row from §3.1.2 Technical Exceptions Matrix
    sap_phrases = [
        "Instance name was not found in the SAP logon",
        "SAP logon",
        "Vx RPA community",
        "Login into SAP",
        "task fails to log into SAP",
    ]
    for phrase in sap_phrases:
        deleted = delete_table_row_containing(doc, phrase)
        if deleted:
            changes.append(f"Deleted SAP exception row (matched: '{phrase}')")
            break

    # Also delete any standalone paragraphs containing SAP RPA text
    for para in list(doc.paragraphs):
        if any(phrase in para.text for phrase in sap_phrases):
            delete_paragraph(para)
            changes.append("Deleted SAP paragraph")

    # (b) Fix "Microsoft Graph" → specific connectors note
    n = replace_in_doc(doc,
        "Microsoft Graph",
        "Office 365 (Outlook/Users/Groups), Teams, Planner, SharePoint, Power BI, Dataverse and Azure AD")
    if n:
        changes.append(f"'Microsoft Graph' → specific connectors ×{n}")

    save_doc(doc, path, "BCP.docx")
    print(f"    changes: {', '.join(changes) if changes else 'none'}")

# ── 7. Informal Testing Log.xlsx ─────────────────────────────────────────────
def fix_testing_log():
    path = FOLDER / "My Meetings Twin - Informal Testing Log.xlsx"
    wb = openpyxl.load_workbook(path)
    changes = []

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    # (a) Typo fix
                    if "Copilto Studio" in cell.value:
                        cell.value = cell.value.replace("Copilto Studio", "Copilot Studio")
                        changes.append(f"Typo 'Copilto' fixed at [{ws.title}]!{cell.coordinate}")
                    # (b) Section header: "Initiate DIA" → "Manage DIA (Red-KPI detection)"
                    if cell.value.strip() == "Initiate DIA":
                        cell.value = "Manage DIA (Red-KPI Detection)"
                        changes.append(f"Section header fixed at [{ws.title}]!{cell.coordinate}")
                    # (c) Section header: "Generate Meeting Deck" → "Post-Meeting"
                    if "Generate Meeting Deck" in cell.value:
                        cell.value = cell.value.replace("Generate Meeting Deck", "Post-Meeting")
                        changes.append(f"'Generate Meeting Deck' → 'Post-Meeting' at [{ws.title}]!{cell.coordinate}")
                    # (d) Tech owner: Dipendra Shekhawat → Vinciane Polet
                    if "Dipendra Shekhawat" in cell.value and "back" not in cell.value.lower():
                        cell.value = cell.value.replace("Dipendra Shekhawat", "Vinciane Polet")
                        changes.append(f"Tech owner aligned at [{ws.title}]!{cell.coordinate}")
                    # (e) Test S.No 61 reword
                    if "DIA process initiates automatically" in cell.value:
                        cell.value = cell.value.replace(
                            "DIA process initiates automatically",
                            "Scheduled flow (PA_SendDIANotificationToOrganizer) reminds organizer; DIA process begins only when organizer opens Manage DIA and confirms")
                        changes.append(f"S.No 61 reworded at [{ws.title}]!{cell.coordinate}")
                    # (f) Mark open failures with a note
                    if "Verify error handling logs failures" in cell.value:
                        # find the result cell in same row and note it
                        pass  # handled by keeping F - just leave as-is, noted in audit

    wb.save(path)
    print(f"  [saved] Informal Testing Log.xlsx")
    print(f"    changes: {', '.join(changes) if changes else 'none'}")

# ── 8. User Guidebook V1.docx ────────────────────────────────────────────────
def fix_guidebook():
    path = FOLDER / "My_Meetings_Twin_Guidebook_V1.docx"
    doc = Document(path)
    changes = []

    # (a) Follow-up meeting: "suggests" → "creates and sends invite"
    suggest_pairs = [
        ("Suggests a follow-up meeting when the transcript indicates unresolved items",
         "Creates a follow-up meeting and sends invites when the transcript indicates unresolved items (on organiser confirmation via CreateCalendarMeeting flow)"),
        ("suggests a follow-up meeting when the transcript indicates unresolved items",
         "creates a follow-up meeting and sends invites when the transcript indicates unresolved items (on organiser confirmation via CreateCalendarMeeting flow)"),
        ("follow-up meeting suggestions",
         "follow-up meeting creation (agent creates the calendar event and sends invites on confirmation)"),
        ("suggest follow-up meetings where required",
         "create follow-up meetings and send invites where required (on organiser confirmation)"),
        ("Suggests a follow-up meeting",
         "Creates a follow-up meeting and sends invites (on organiser confirmation)"),
        ("suggests a follow-up meeting",
         "creates a follow-up meeting and sends invites (on organiser confirmation)"),
    ]
    for old_s, new_s in suggest_pairs:
        n = replace_in_doc(doc, old_s, new_s)
        if n:
            changes.append(f"Follow-up meeting wording fixed ×{n}")

    # (b) Fix DIA form fields in §3.2
    field_pairs = [
        # The audit says the Guidebook lists KPI context fields instead of real DIA capture fields
        ("executive summary, key updates, next steps, RAG context, and other DIA fields",
         "Situation, Current Response, Patient Impact, Balanced Scorecard, DIA, Recommendation, Next Steps, and DIA Type (Decision / Input / Awareness)"),
        ("executive summary, key updates, next steps, RAG context",
         "Situation, Current Response, Patient Impact, Balanced Scorecard, DIA, Recommendation, Next Steps, and DIA Type (Decision / Input / Awareness)"),
    ]
    for old_f, new_f in field_pairs:
        n = replace_in_doc(doc, old_f, new_f)
        if n:
            changes.append(f"DIA form fields corrected ×{n}")

    # (c) Add Meeting Purpose step note
    n = replace_in_doc(doc,
        "Tier ID → KPI/MS Project",
        "Tier ID → Meeting Purpose (free text, used in the agenda) → KPI/MS Project")
    if n:
        changes.append("Meeting Purpose step added to walkthrough")

    # (d) Notifications section: organiser-vs-contributor correction
    n = replace_in_doc(doc,
        "KPI owners and attendees receive prompts before inputs are due",
        "the organiser receives a preparation reminder (PA_SendDIANotificationToOrganizer); owners and attendees are notified only after the organiser opens Manage DIA, reviews flagged items and confirms — which creates the DIA tasks and triggers the owner notification (Teams + email)")
    if n:
        changes.append("Notifications section corrected (organiser vs contributor)")

    save_doc(doc, path, "User Guidebook V1.docx")
    print(f"    changes: {', '.join(changes) if changes else 'none'}")

# ── main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("=== Backing up files ===")
    backup_folder()

    print("\n=== 1. TDD.docx ===")
    fix_tdd()

    print("\n=== 2. AI Accountability Report.xlsx ===")
    fix_ai_accountability()

    print("\n=== 3. User Requirements.docx ===")
    fix_user_requirements()

    print("\n=== 4. EUD Checklist.docx ===")
    fix_eud_checklist()

    print("\n=== 5. GSC RAI Process.docx ===")
    fix_gsc_rai_process()

    print("\n=== 6. BCP.docx ===")
    fix_bcp()

    print("\n=== 7. Informal Testing Log.xlsx ===")
    fix_testing_log()

    print("\n=== 8. User Guidebook V1.docx ===")
    fix_guidebook()

    print("\n=== All done ===")
